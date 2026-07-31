"""Collect bounded observations from one exact Phase 1B candidate image."""

from __future__ import annotations

import hashlib
import importlib.metadata
import importlib.util
import json
import os
import re
import shutil
import stat
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Callable


MANIFEST_PATH = Path("/opt/gdpval/v2/capabilities.json")
MAX_OUTPUT_BYTES = 65536
COMMAND_TIMEOUT_SECONDS = 45
_HEX_SHA256 = re.compile(r"[0-9a-f]{64}\Z", re.ASCII)


def _canonical_sha256(value) -> str:
    return hashlib.sha256(
        json.dumps(
            value,
            sort_keys=True,
            separators=(",", ":"),
            ensure_ascii=True,
            allow_nan=False,
        ).encode("utf-8")
    ).hexdigest()


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    total = 0
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            total += len(chunk)
            if total > 1024 * 1024 * 1024:
                raise RuntimeError(f"probe file exceeds 1 GiB: {path}")
            digest.update(chunk)
    return digest.hexdigest()


def _run(
    argv: list[str],
    *,
    cwd: Path,
    timeout: int = COMMAND_TIMEOUT_SECONDS,
) -> subprocess.CompletedProcess:
    command = list(argv)
    if Path(command[0]).name in {"python", "python3", "python3.11"}:
        command[1:1] = ["-I", "-S", "-B"]
    environment = {
        "HOME": str(cwd / ".home"),
        "TMPDIR": str(cwd / ".tmp"),
        "XDG_CACHE_HOME": str(cwd / ".cache"),
        "XDG_CONFIG_HOME": str(cwd / ".config"),
        "LANG": "C.UTF-8",
        "LC_ALL": "C.UTF-8",
        "PATH": "/usr/local/bin:/usr/bin:/bin",
        "PYTHONDONTWRITEBYTECODE": "1",
    }
    for name in (".home", ".tmp", ".cache", ".config"):
        (cwd / name).mkdir(mode=0o700, parents=True, exist_ok=True)
    completed = subprocess.run(
        command,
        cwd=cwd,
        env=environment,
        stdin=subprocess.DEVNULL,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        timeout=timeout,
        check=False,
    )
    if len(completed.stdout) > MAX_OUTPUT_BYTES or len(completed.stderr) > MAX_OUTPUT_BYTES:
        raise RuntimeError(f"probe output exceeded limit: {argv[0]}")
    if completed.returncode != 0:
        stderr = completed.stderr.decode("utf-8", errors="replace")[:500]
        raise RuntimeError(
            f"probe command failed ({completed.returncode}): {argv[0]}: {stderr}"
        )
    return completed


def _command_observation(item: dict, work: Path) -> dict:
    path_value = shutil.which(item["name"])
    if not path_value:
        raise RuntimeError(f"required command is missing: {item['name']}")
    path = Path(path_value).resolve(strict=True)
    completed = _run(item["probe"], cwd=work)
    output = (completed.stdout + b"\n" + completed.stderr).decode(
        "utf-8", errors="replace"
    ).strip()
    if not output:
        raise RuntimeError(f"required command returned no version: {item['name']}")
    return {
        "name": item["name"],
        "version": output.splitlines()[0][:500],
        "sha256": _file_sha256(path),
    }


def _distribution_for_module(name: str):
    mapping = importlib.metadata.packages_distributions()
    candidates = mapping.get(name, [])
    aliases = {
        "PIL": "Pillow",
        "pptx": "python-pptx",
        "sklearn": "scikit-learn",
    }
    if aliases.get(name):
        candidates = [aliases[name], *candidates]
    for candidate in candidates:
        try:
            return importlib.metadata.distribution(candidate)
        except importlib.metadata.PackageNotFoundError:
            continue
    raise RuntimeError(f"distribution is missing for module: {name}")


def _distribution_sha256(distribution) -> str:
    records = []
    for item in sorted(distribution.files or [], key=str):
        path = Path(distribution.locate_file(item))
        if not path.is_file() or path.is_symlink():
            continue
        records.append({
            "path": str(item),
            "sha256": _file_sha256(path),
            "size": path.stat().st_size,
        })
    if not records:
        raise RuntimeError(f"distribution has no files: {distribution.metadata['Name']}")
    return _canonical_sha256(records)


def _module_observation(item: dict) -> dict:
    spec = importlib.util.find_spec(item["name"])
    if spec is None:
        raise RuntimeError(f"required Python module is missing: {item['name']}")
    distribution = _distribution_for_module(item["name"])
    return {
        "name": item["name"],
        "version": distribution.version,
        "sha256": _distribution_sha256(distribution),
    }


def _font_observation(name: str, work: Path) -> dict:
    completed = _run(
        ["fc-match", "--format", "%{family}|%{file}", name],
        cwd=work,
    )
    value = completed.stdout.decode("utf-8", errors="strict")
    family, separator, path_value = value.partition("|")
    if not separator or name not in family.split(","):
        raise RuntimeError(f"required font family is missing: {name}")
    path = Path(path_value).resolve(strict=True)
    return {
        "name": name,
        "version": path.name,
        "sha256": _file_sha256(path),
    }


def _effective_sbom_namespace(expected_sha256: str) -> dict:
    if _HEX_SHA256.fullmatch(expected_sha256) is None:
        raise RuntimeError("effective SBOM inventory digest is invalid")
    path = Path(__file__).with_name("effective_sbom.py")
    required_flags = ("O_NOFOLLOW", "O_CLOEXEC")
    if any(not isinstance(getattr(os, name, None), int) for name in required_flags):
        raise RuntimeError("secure Unix open flags are required for image inventory")
    descriptor = os.open(path, os.O_RDONLY | os.O_NOFOLLOW | os.O_CLOEXEC)
    try:
        before = os.fstat(descriptor)
        if (
            not stat.S_ISREG(before.st_mode)
            or before.st_nlink != 1
            or before.st_size < 0
            or before.st_size > 1024 * 1024
        ):
            raise RuntimeError("effective SBOM inventory module is invalid")
        chunks = []
        while True:
            chunk = os.read(descriptor, 64 * 1024)
            if not chunk:
                break
            chunks.append(chunk)
        after = os.fstat(descriptor)
    finally:
        os.close(descriptor)
    source = b"".join(chunks)
    def identity(value):
        return (
            value.st_dev, value.st_ino, value.st_mode, value.st_nlink,
            value.st_size, value.st_mtime_ns,
        )

    if identity(before) != identity(after) or len(source) != before.st_size:
        raise RuntimeError("effective SBOM inventory module changed while reading")
    if hashlib.sha256(source).hexdigest() != expected_sha256:
        raise RuntimeError("effective SBOM inventory module digest differs")
    namespace = {
        "__builtins__": __builtins__,
        "__file__": str(path),
        "__name__": "_gdpval_agentic_v2_effective_sbom",
        "__package__": None,
    }
    exec(compile(source, "effective_sbom.py", "exec"), namespace)
    if not callable(namespace.get("r_inventory_records")):
        raise RuntimeError("effective SBOM R inventory is unavailable")
    return namespace


def _r_inventory_records(expected_sbom_sha256: str) -> list[str]:
    records = _effective_sbom_namespace(expected_sbom_sha256)[
        "r_inventory_records"
    ]()
    values = sorted(f"{item['name']}={item['version']}" for item in records)
    if len(values) != len(set(values)):
        raise RuntimeError("R package observation is duplicated")
    return values


def _inventory(work: Path, expected_sbom_sha256: str) -> dict:
    debian = _run(
        ["dpkg-query", "-W", "-f=${Package}:${Architecture}=${Version}\\n"], cwd=work
    ).stdout.decode("utf-8").splitlines()
    python = sorted({
        f"{distribution.metadata.get('Name') or distribution.name}={distribution.version}"
        for distribution in importlib.metadata.distributions()
    })
    r_packages = _r_inventory_records(expected_sbom_sha256)
    npm_output = _run(["npm", "--version"], cwd=work).stdout.decode("utf-8").strip()
    records = {
        "debian": sorted(debian),
        "python": python,
        "r": r_packages,
        "npm": [f"npm={npm_output}"],
    }
    return {
        name: {
            "count": len(values),
            "sha256": _canonical_sha256(values),
            "records": values,
        }
        for name, values in records.items()
    }


def _hash_artifacts(paths: list[Path]) -> str:
    return _canonical_sha256([
        {
            "name": path.name,
            "sha256": _file_sha256(path),
            "size": path.stat().st_size,
        }
        for path in sorted(paths, key=lambda value: value.name)
    ])


def _semantic_artifact(work: Path, name: str, value) -> str:
    path = work / name
    path.write_text(
        json.dumps(value, sort_keys=True, separators=(",", ":")),
        encoding="utf-8",
    )
    return _hash_artifacts([path])


def _smoke_browser(work: Path) -> str:
    html = work / "browser.html"
    screenshot = work / "browser.png"
    html.write_text(
        "<!doctype html><meta charset='utf-8'><title>GDPVal</title>"
        "<h1 style='color:#123456'>GDPVAL Phase 1B</h1>",
        encoding="utf-8",
    )
    _run([
        "chromium",
        "--headless",
        "--no-sandbox",
        "--disable-gpu",
        "--disable-dev-shm-usage",
        f"--user-data-dir={work / 'chromium'}",
        "--window-size=800,600",
        f"--screenshot={screenshot}",
        html.as_uri(),
    ], cwd=work)
    if screenshot.stat().st_size < 1000:
        raise RuntimeError("Chromium screenshot is unexpectedly small")
    return _hash_artifacts([html, screenshot])


def _smoke_compilers(work: Path) -> str:
    sources = {
        "main.c": "#include <stdio.h>\nint main(){puts(\"C-OK\");return 0;}\n",
        "main.cpp": "#include <iostream>\nint main(){std::cout<<\"CPP-OK\\n\";}\n",
        "main.f90": "program p\nprint *, 'F-OK'\nend program p\n",
    }
    outputs = []
    for name, content in sources.items():
        path = work / name
        path.write_text(content, encoding="utf-8")
        outputs.append(path)
    for compiler, source, output in (
        ("gcc", "main.c", "c.out"),
        ("g++", "main.cpp", "cpp.out"),
        ("gfortran", "main.f90", "fortran.out"),
    ):
        _run([compiler, source, "-o", output], cwd=work)
        completed = _run([str(work / output)], cwd=work)
        if b"OK" not in completed.stdout:
            raise RuntimeError(f"compiled smoke did not execute: {compiler}")
        outputs.append(work / output)
    cmake_source = work / "cmake-src"
    cmake_build = work / "cmake-build"
    cmake_source.mkdir()
    (cmake_source / "CMakeLists.txt").write_text(
        "cmake_minimum_required(VERSION 3.20)\nproject(gdpval C)\n"
        "add_executable(gdpval ../main.c)\n",
        encoding="utf-8",
    )
    _run(["cmake", "-S", str(cmake_source), "-B", str(cmake_build)], cwd=work)
    _run(["cmake", "--build", str(cmake_build)], cwd=work)
    outputs.append(cmake_build / "gdpval")
    return _hash_artifacts(outputs)


def _smoke_ml(work: Path) -> str:
    import numpy as np
    from sklearn.linear_model import LinearRegression

    x = np.array([[0.0], [1.0], [2.0], [3.0]])
    y = np.array([1.0, 3.0, 5.0, 7.0])
    model = LinearRegression().fit(x, y)
    payload = {
        "coefficient": float(model.coef_[0]),
        "intercept": float(model.intercept_),
        "score": float(model.score(x, y)),
    }
    path = work / "ml.json"
    path.write_text(json.dumps(payload, sort_keys=True), encoding="utf-8")
    return _hash_artifacts([path])


def _smoke_dxf(work: Path) -> str:
    import ezdxf

    path = work / "drawing.dxf"
    document = ezdxf.new("R2010")
    document.modelspace().add_line((0, 0), (10, 5))
    document.saveas(path)
    loaded = ezdxf.readfile(path)
    lines = list(loaded.modelspace().query("LINE"))
    if len(lines) != 1:
        raise RuntimeError("DXF roundtrip lost geometry")
    line = lines[0]
    return _semantic_artifact(work, "dxf-semantic.json", {
        "entity_count": len(lines),
        "start": [float(value) for value in line.dxf.start],
        "end": [float(value) for value in line.dxf.end],
        "version": loaded.dxfversion,
    })


def _smoke_gis(work: Path) -> str:
    import geopandas as gpd
    from shapely.geometry import Point

    path = work / "points.gpkg"
    frame = gpd.GeoDataFrame(
        {"name": ["alpha", "beta"]},
        geometry=[Point(0, 0), Point(1, 1)],
        crs="EPSG:4326",
    )
    frame.to_file(path, driver="GPKG")
    loaded = gpd.read_file(path)
    if len(loaded) != 2 or loaded.crs is None:
        raise RuntimeError("GeoPackage roundtrip failed")
    return _semantic_artifact(work, "gis-semantic.json", {
        "columns": list(loaded.columns),
        "crs": loaded.crs.to_string(),
        "rows": [
            {"name": row["name"], "geometry": row.geometry.wkt}
            for _, row in loaded.sort_values("name").iterrows()
        ],
    })


def _smoke_media(work: Path) -> str:
    video = work / "video.mkv"
    audio = work / "audio.wav"
    _run([
        "ffmpeg", "-hide_banner", "-loglevel", "error", "-y",
        "-f", "lavfi", "-i", "testsrc=size=64x64:rate=1",
        "-t", "1", "-c:v", "ffv1", str(video),
    ], cwd=work)
    _run([
        "ffmpeg", "-hide_banner", "-loglevel", "error", "-y",
        "-f", "lavfi", "-i", "sine=frequency=440:duration=1",
        str(audio),
    ], cwd=work)
    summaries = []
    for path in (video, audio):
        completed = _run([
            "ffprobe", "-v", "error", "-show_streams", "-of", "json", str(path)
        ], cwd=work)
        document = json.loads(completed.stdout)
        if not document.get("streams"):
            raise RuntimeError("ffprobe returned no streams")
        summaries.append({
            "name": path.name,
            "streams": [{
                key: stream[key]
                for key in (
                    "codec_name", "codec_type", "width", "height",
                    "sample_rate", "channels", "r_frame_rate",
                )
                if key in stream
            } for stream in document["streams"]],
        })
    return _semantic_artifact(work, "media-semantic.json", summaries)


def _smoke_ocr(work: Path) -> str:
    from PIL import Image, ImageDraw, ImageFont

    image_path = work / "ocr.png"
    output_path = work / "ocr.txt"
    font = ImageFont.truetype(
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 64
    )
    image = Image.new("RGB", (800, 180), "white")
    ImageDraw.Draw(image).text((20, 40), "GDPVAL 42", fill="black", font=font)
    image.save(image_path)
    completed = _run([
        "tesseract", str(image_path), "stdout", "--psm", "7", "-l", "eng"
    ], cwd=work)
    text = completed.stdout.decode("utf-8", errors="replace").strip()
    if "GDPVAL" not in text.upper():
        raise RuntimeError(f"OCR did not recover marker: {text[:100]}")
    output_path.write_text(text, encoding="utf-8")
    return _hash_artifacts([image_path, output_path])


def _smoke_office(work: Path) -> str:
    from docx import Document

    docx_path = work / "office.docx"
    document = Document()
    document.add_heading("GDPVal Phase 1B", level=1)
    document.add_paragraph("Professional-work substrate office roundtrip.")
    document.save(docx_path)
    _run([
        "libreoffice", "--headless", "--convert-to", "pdf", "--outdir",
        str(work), str(docx_path),
    ], cwd=work, timeout=90)
    pdf_path = work / "office.pdf"
    text_path = work / "office.txt"
    _run(["pdftotext", str(pdf_path), str(text_path)], cwd=work)
    text = text_path.read_text(encoding="utf-8")
    if "GDPVal Phase 1B" not in text:
        raise RuntimeError("Office/PDF roundtrip lost marker")
    normalized = " ".join(text.split())
    return _semantic_artifact(work, "office-semantic.json", {
        "marker": "GDPVal Phase 1B",
        "text": normalized,
    })


def _smoke_spreadsheet(work: Path) -> str:
    from openpyxl import Workbook, load_workbook

    path = work / "formula.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet["A1"] = 2
    sheet["A2"] = 3
    sheet["A3"] = "=SUM(A1:A2)"
    workbook.save(path)
    loaded = load_workbook(path, data_only=False)
    if loaded.active["A3"].value != "=SUM(A1:A2)":
        raise RuntimeError("spreadsheet formula roundtrip failed")
    return _semantic_artifact(work, "spreadsheet-semantic.json", {
        "A1": loaded.active["A1"].value,
        "A2": loaded.active["A2"].value,
        "A3": loaded.active["A3"].value,
        "sheet": loaded.active.title,
    })


SMOKES: dict[str, Callable[[Path], str]] = {
    "browser-local-screenshot": _smoke_browser,
    "compiler-matrix": _smoke_compilers,
    "data-ml-fit": _smoke_ml,
    "dxf-roundtrip": _smoke_dxf,
    "gis-geopackage": _smoke_gis,
    "media-generate-probe": _smoke_media,
    "ocr-local-image": _smoke_ocr,
    "office-pdf-roundtrip": _smoke_office,
    "spreadsheet-formula-roundtrip": _smoke_spreadsheet,
}


def collect(expected_sbom_sha256: str) -> dict:
    if os.geteuid() != 65532 or os.getegid() != 65532:
        raise RuntimeError("Phase 1B probe must run as UID/GID 65532")
    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
    manifest_sha256 = _canonical_sha256(manifest)
    with tempfile.TemporaryDirectory(prefix="phase1b-", dir="/work") as temporary:
        root = Path(temporary)
        commands = [
            _command_observation(item, root)
            for item in manifest["commands"]
        ]
        modules = [
            _module_observation(item)
            for item in manifest["python_modules"]
        ]
        fonts = [
            _font_observation(name, root)
            for name in manifest["font_families"]
        ]
        smokes = []
        for item in manifest["smoke_matrix"]:
            smoke_root = root / item["id"]
            smoke_root.mkdir(mode=0o700)
            smokes.append({
                "id": item["id"],
                "status": "pass",
                "artifact_sha256": SMOKES[item["id"]](smoke_root),
            })
        inventory = _inventory(root, expected_sbom_sha256)
    return {
        "schema_version": "1.0",
        "substrate_id": "professional-work-v1",
        "manifest_sha256": manifest_sha256,
        "commands": commands,
        "python_modules": modules,
        "font_families": fonts,
        "smokes": smokes,
        "package_inventory": inventory,
    }


def main() -> None:
    if len(sys.argv) != 2:
        raise RuntimeError("effective SBOM inventory digest argument is required")
    print(json.dumps(collect(sys.argv[1]), sort_keys=True, separators=(",", ":")))


if __name__ == "__main__":
    main()