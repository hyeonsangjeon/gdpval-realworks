"""A filter that selects nothing is a fact about the filter, not about the file.

``read_content`` on a workbook took ``scope={"sheet": ...}``. A name that
matched no sheet selected no sheets, produced an empty string, and the empty
string was returned as::

    {"text": "", "char_count": 0, "has_text_layer": false,
     "note": "no extractable text -- an empty text read means this file
              carries no extractable text, NOT that the content is absent"}

Every word after ``note`` is about the file. The file was never asked. Matching
is exact and case-sensitive, so ``"prepaid summary"`` misses ``"Prepaid
Summary"`` and the workbook is declared textless.

**Measured, in the 185-task gold-ceiling run** (payload
``cef3a5b9…/exp_gold_baseline__judge_gpt-5_6-sol__gold_ceiling_185_v2_sol_max``):
15 rubric items across 9 tasks were graded with that sentence as their entire
``evidence``. All 15 scored ``fail`` at 0.0 against 29.0 points, all routed
``text``, each after three ``read_deliverable`` calls. Reading all 10 files
they name with ``scope={}`` — no model, no judge — returns 346 to 200,000
characters of text from every one, and none is missing from disk. The four
30-task Stage 1 payloads, graded before the disclaimer existed, carry it zero
times, and two of these same workbooks were read successfully there.

``320-three-gaps-that-closed.md`` counts 23 where this counts 15, and both are
right about what they counted. The judge did not always quote the note whole,
so the needle decides the total: the whole sentence matches 17 items, the tail
``"NOT that the content is absent"`` matches 23, the union 24. Restricted to
office files every rule returns the same 15 and the same 0.0 of 29.0 — the
entire spread is nine rubric items on two ``.mp4`` deliverables
(``GreenEnergy_v1.mp4``, ``Goodsin Studios CG Reel 2025.mp4``), which are a
routing question and not this one. Both facts are pinned below.

The 15 cannot be attributed further than that: ``judge_raw_response`` is null
and ``tools_used`` records the three calls by name without their arguments, so
the payload preserves *that* the tool was called and not *what was asked*. The
mechanism below reproduces model-free; the per-item attribution does not, and
this file does not claim it.

Two of the nine tasks are the ones ``320`` left open. It observed that
``7d7fc9a7`` and ``dfb4e0cd`` read fine in Stage 1 and came back empty in
Stage 3, and recorded that it had not found out why. Those two files are
``Aurisic_Amortization_4-25.xlsx`` and
``Time Elapsed vs Funds Spent Analysis 2025.03.31.xlsx`` — 12,399 and 200,000
characters, both still readable today, and the first is the workbook the
before/after matrix in the pull request was measured on.

The tool already refuses this shape three times over — ``no member 'x' in
archive; members: [...]``, ``page 999 out of range 1..3``, ``slide 999 out of
range 1..10`` — and refuses unknown keys in ``render_to_image``. The read and
formatting ops were the gap. This file holds them to the same behaviour.

Nothing here calls a model, grades anything, or spends anything.
"""

from __future__ import annotations

import json
import os
from pathlib import Path

import pytest

from core.tools.read_deliverable import (
    MODEL_READ_DELIVERABLE_TOOL_SCHEMA,
    READ_DELIVERABLE_TOOL_SCHEMA,
    _EMPTY_READ_DISCLAIMER,
    read_deliverable,
)

REPO_ROOT = Path(__file__).resolve().parents[2]
GRADES_ROOT = REPO_ROOT / "data/grades/_diagnostic"

#: Exactly what the paid payload holds. See the module docstring.
MEASURED_ITEMS = 15
MEASURED_TASKS = 9
MEASURED_MAX_SCORE = 29.0
MEDIA_SUFFIXES = {".mp4", ".mov", ".avi", ".mkv", ".wav", ".mp3", ".m4a",
                  ".png", ".jpg", ".jpeg"}

#: Three ways to count the same sentence in the same payload. The judge did
#: not always quote it whole, so the substring chosen changes the total:
#: the whole sentence matches 17 items, the tail alone matches 23 (which is
#: the figure `320-three-gaps-that-closed.md` published), and the union is 24.
#: On office files all three give 15 -- the entire spread is two ``.mp4``
#: deliverables. That stability is the point of pinning all three.
_FULL_SENTENCE = _EMPTY_READ_DISCLAIMER[:60]
_TAIL_FRAGMENT = "NOT that the content is absent"

SHEETS = ("Prepaid Summary", "PPD Exp #1250", "PPD Ins #1251")


@pytest.fixture
def base_dir(tmp_path: Path) -> Path:
    return tmp_path


@pytest.fixture
def workbook(base_dir: Path) -> Path:
    """Three named sheets, all of them holding text. Shaped after
    ``Aurisic_Amortization_4-25.xlsx``, one of the 10 measured files."""
    openpyxl = pytest.importorskip("openpyxl")
    wb = openpyxl.Workbook()
    wb.active.title = SHEETS[0]
    for name in SHEETS[1:]:
        wb.create_sheet(name)
    for name in SHEETS:
        ws = wb[name]
        ws["A1"] = f"{name} header"
        ws["A2"] = "balance"
        ws["B2"] = 1250
    p = base_dir / "amortization.xlsx"
    wb.save(p)
    return p


@pytest.fixture
def three_page_pdf(base_dir: Path) -> Path:
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas

    p = base_dir / "report.pdf"
    c = canvas.Canvas(str(p))
    for n in (1, 2, 3):
        c.drawString(100, 750, f"Page {n} carries real text.")
        c.showPage()
    c.save()
    return p


@pytest.fixture
def memo(base_dir: Path) -> Path:
    pytest.importorskip("docx")
    from docx import Document

    doc = Document()
    doc.add_paragraph("Body text that read_content returns in full.")
    p = base_dir / "memo.docx"
    doc.save(p)
    return p


def _call(op, path, base_dir, scope=None):
    return read_deliverable(op, path.name, base_dir=str(base_dir), scope=scope)


def _refused(response) -> str:
    """The error text, asserting the envelope refused rather than answered.

    ``bad_scope`` matters as much as the refusal: it is the retryable kind the
    judge already handles, so a caller that named the wrong sheet gets a turn
    to name the right one instead of a dead end."""
    assert response["ok"] is False, f"expected a refusal, got {response}"
    assert response["error_type"] == "bad_scope", response
    return response["error"]


# ── 1. the defect: a sheet that is not in the workbook ────────────────


def test_a_sheet_that_does_not_exist_is_refused_not_reported_as_no_text(
    base_dir, workbook
):
    message = _refused(
        _call("read_content", workbook, base_dir, {"sheet": "Sheet1"})
    )
    assert "no sheet 'Sheet1' in workbook" in message
    for name in SHEETS:
        assert name in message, "the refusal must name the sheets that exist"


def test_the_refusal_never_carries_the_empty_read_disclaimer(base_dir, workbook):
    """The whole point. The sentence that graded 15 items to zero must not be
    reachable from a scope that selected nothing."""
    for scope in ({"sheet": "Sheet1"}, {"sheet": "Amortization"},
                  {"sheet": "1250"}, {"sheet": ""}):
        blob = json.dumps(_call("read_content", workbook, base_dir, scope))
        assert _EMPTY_READ_DISCLAIMER not in blob, scope
        assert "no extractable text" not in blob, scope


def test_a_sheet_name_in_the_wrong_case_is_refused_with_the_right_case(
    base_dir, workbook
):
    """Matching is exact, which is defensible. Silently reporting the miss as
    an empty file is not — and the message has to carry the fix."""
    message = _refused(
        _call("read_content", workbook, base_dir, {"sheet": "prepaid summary"})
    )
    assert "no sheet 'prepaid summary' in workbook" in message
    assert "Prepaid Summary" in message


def test_a_sheet_that_does_exist_still_reads(base_dir, workbook):
    response = _call("read_content", workbook, base_dir, {"sheet": SHEETS[1]})
    assert response["ok"] is True
    assert SHEETS[1] in response["data"]["text"]
    assert SHEETS[0] not in response["data"]["text"], "the narrowing must apply"


def test_no_scope_still_reads_every_sheet(base_dir, workbook):
    response = _call("read_content", workbook, base_dir)
    assert response["ok"] is True
    for name in SHEETS:
        assert name in response["data"]["text"]


# ── 2. the same defect in inspect_formatting ─────────────────────────


def test_inspect_formatting_refuses_an_absent_sheet_rather_than_returning_none(
    base_dir, workbook
):
    """This half was worse: it returned ``{"sheets": []}`` with no note at all.
    A bare empty list is the "absence is not zero" failure with nothing
    attached to argue against it."""
    message = _refused(
        _call("inspect_formatting", workbook, base_dir, {"sheet": "Sheet1"})
    )
    assert "no sheet 'Sheet1' in workbook" in message
    assert SHEETS[0] in message


def test_inspect_formatting_with_a_real_sheet_returns_that_sheet(
    base_dir, workbook
):
    response = _call("inspect_formatting", workbook, base_dir,
                     {"sheet": SHEETS[2]})
    assert response["ok"] is True
    assert [s["name"] for s in response["data"]["sheets"]] == [SHEETS[2]]


# ── 3. the same defect on a PDF page window ──────────────────────────


def test_a_page_window_past_the_end_is_refused_not_called_a_missing_text_layer(
    base_dir, three_page_pdf
):
    """``page_start=53`` on a 3-page PDF used to answer "this PDF has no text
    layer" — about a document whose first page yields real text."""
    message = _refused(
        _call("read_content", three_page_pdf, base_dir,
              {"page_start": 53, "page_end": 62})
    )
    assert "page_start 53 out of range 1..3" in message
    assert "no text layer" not in message


def test_a_backwards_page_window_is_refused(base_dir, three_page_pdf):
    message = _refused(
        _call("read_content", three_page_pdf, base_dir,
              {"page_start": 3, "page_end": 1})
    )
    assert "page_end 1 is before page_start 3" in message


def test_a_page_window_inside_the_document_still_reads(base_dir, three_page_pdf):
    response = _call("read_content", three_page_pdf, base_dir,
                     {"page_start": 2, "page_end": 2})
    assert response["ok"] is True
    assert "Page 2" in response["data"]["text"]
    assert "Page 1 carries" not in response["data"]["text"]


def test_a_page_number_that_is_not_a_number_is_a_scope_error(
    base_dir, three_page_pdf
):
    message = _refused(
        _call("read_content", three_page_pdf, base_dir, {"page_start": "first"})
    )
    assert "page_start must be a positive 1-based integer" in message


# ── 4. a key the op does not implement ───────────────────────────────


@pytest.mark.parametrize(
    "fixture_name, scope, unknown",
    [
        ("three_page_pdf", {"sheet": "Summary"}, "sheet"),
        ("memo", {"page": 999}, "page"),
        ("memo", {"sheet": "Summary"}, "sheet"),
        ("workbook", {"page_start": 1}, "page_start"),
    ],
)
def test_a_key_this_op_does_not_implement_is_refused_not_dropped(
    request, base_dir, fixture_name, scope, unknown
):
    """Dropping it returns the whole file to a caller who believes it asked
    for a part of it, and evidence then gets quoted against a page or sheet
    that was never selected. ``render_to_image`` has always refused this."""
    target = request.getfixturevalue(fixture_name)
    message = _refused(_call("read_content", target, base_dir, scope))
    assert f"'{unknown}'" in message
    assert "allowed keys" in message


def test_the_refusal_names_the_keys_that_would_have_worked(
    base_dir, three_page_pdf
):
    message = _refused(
        _call("read_content", three_page_pdf, base_dir, {"page": 1})
    )
    assert "page_start" in message and "page_end" in message


def test_inspect_formatting_refuses_a_key_it_does_not_implement(
    base_dir, workbook
):
    """The same gap on the other op. ``inspect_formatting`` reads only
    ``sheet``; a page window handed to it described nothing and changed
    nothing, and the full-workbook answer came back as if it had."""
    message = _refused(
        _call("inspect_formatting", workbook, base_dir,
              {"sheet": SHEETS[0], "page_start": 1})
    )
    assert "inspect_formatting on xlsx" in message
    assert "'page_start'" in message
    assert "allowed keys: ['sheet']" in message


# ── 5. the disclaimer must survive where it is true ───────────────────


def test_a_document_that_really_holds_no_text_still_says_so(base_dir):
    """A ``.docx`` with nothing in it is the closest true neighbour of the
    false claim. The 15 items include one ``.docx``, and its 346 characters
    are the difference between this response and that one — so both have to
    stay reachable, each from its own cause."""
    pytest.importorskip("docx")
    from docx import Document

    p = base_dir / "empty.docx"
    Document().save(p)

    response = _call("read_content", p, base_dir)
    assert response["ok"] is True
    assert response["data"]["char_count"] == 0
    assert response["data"]["has_text_layer"] is False
    assert _EMPTY_READ_DISCLAIMER in response["data"]["note"]


def test_a_pdf_that_really_has_no_text_layer_still_says_so(base_dir):
    """The other half of the negative control, and the one that matters most
    here: ``page_start`` past the end used to borrow *this* sentence. Taking
    it away from the out-of-range case must not take it away from the scanned
    document it was written for."""
    pytest.importorskip("reportlab")
    from reportlab.pdfgen import canvas

    p = base_dir / "scanned.pdf"
    c = canvas.Canvas(str(p))
    c.rect(100, 100, 200, 200, fill=1)  # ink, no glyphs
    c.showPage()
    c.save()

    response = _call("read_content", p, base_dir)
    assert response["ok"] is True
    assert response["data"]["has_text_layer"] is False
    assert _EMPTY_READ_DISCLAIMER in response["data"]["note"]


def test_a_file_that_really_holds_no_text_still_says_so(base_dir):
    """The fix removes a false claim, not the true one. A PNG holds no text
    and must keep saying it — otherwise this change would trade one silent
    failure for another."""
    pytest.importorskip("PIL")
    from PIL import Image

    p = base_dir / "chart.png"
    Image.new("RGB", (16, 16), color="red").save(p)

    response = _call("read_content", p, base_dir)
    assert response["ok"] is True
    assert response["data"]["char_count"] == 0
    assert _EMPTY_READ_DISCLAIMER in response["data"]["note"]


def test_an_empty_but_real_sheet_reads_as_that_sheet_holding_nothing(base_dir):
    """The distinction the fix turns on, from the other side.

    A sheet that exists and holds nothing must still be reachable, and the
    answer must be about *that sheet*: the reader emits its ``[Sheet: Blank]``
    header and no rows, so the response says which sheet was opened and that it
    was bare. Only the sheet that does not exist is refused. If this ever
    started refusing too, the fix would have replaced a false claim about the
    file with a false claim about the request."""
    openpyxl = pytest.importorskip("openpyxl")
    wb = openpyxl.Workbook()
    wb.active.title = "Blank"
    p = base_dir / "blank.xlsx"
    wb.save(p)

    response = _call("read_content", p, base_dir, {"sheet": "Blank"})
    assert response["ok"] is True
    assert response["data"]["text"] == "[Sheet: Blank]\n"
    assert _EMPTY_READ_DISCLAIMER not in json.dumps(response)


# ── 6. the schema names the keys, so nothing has to be guessed ────────


@pytest.mark.parametrize(
    "schema",
    [READ_DELIVERABLE_TOOL_SCHEMA, MODEL_READ_DELIVERABLE_TOOL_SCHEMA],
    ids=["full", "model"],
)
def test_the_schema_names_the_content_scope_keys(schema):
    """Neither schema named one content-scope key, so a judge that wanted a
    single sheet had to guess the key and the value both, and a miss on either
    came back as an empty file."""
    described = schema["parameters"]["properties"]["scope"]["description"]
    assert "sheet" in described
    assert "page_start" in described and "page_end" in described
    assert "refused" in described
    assert "never reported as an empty file" in described


# ── 7. the run this was measured from ────────────────────────────────


def _stage_three_payload():
    for path in sorted(GRADES_ROOT.rglob("*gold_ceiling_185_v2_sol_max*.json")):
        payload = json.loads(path.read_text(encoding="utf-8"))
        if isinstance(payload, dict) and payload.get("tasks"):
            return payload
    return None


def _disclaimer_items(payload, needle=None):
    out = []
    for task in payload.get("tasks") or []:
        for item in task.get("items") or []:
            evidence = item.get("evidence")
            if not isinstance(evidence, str):
                continue
            if needle is None:
                hit = _FULL_SENTENCE in evidence or _TAIL_FRAGMENT in evidence
            else:
                hit = needle in evidence
            if not hit:
                continue
            paths = item.get("selected_paths") or []
            if not paths or any(Path(f).suffix.lower() in MEDIA_SUFFIXES
                                for f in paths):
                continue
            out.append((task, item))
    return out


def test_the_paid_run_this_was_measured_from_still_reads_the_way_it_did():
    """The evidence the fix was built on, pinned. This payload is frozen —
    nothing here regrades it — so a change in these numbers means the payload
    moved, and the argument above needs re-reading before it is trusted."""
    payload = _stage_three_payload()
    if payload is None:
        pytest.skip("the 185-task gold-ceiling payload is not checked in here")

    items = _disclaimer_items(payload, _FULL_SENTENCE)
    assert len(items) == MEASURED_ITEMS
    assert len({t["task_id"] for t, _ in items}) == MEASURED_TASKS
    assert {i["verdict"] for _, i in items} == {"fail"}
    assert sum(float(i["awarded_score"]) for _, i in items) == 0.0
    assert sum(float(i["max_score"]) for _, i in items) == MEASURED_MAX_SCORE
    assert {i["routing_modality"] for _, i in items} == {"text"}


def test_the_office_count_does_not_depend_on_how_the_sentence_is_matched():
    """Why this file says 15 where `320-three-gaps-that-closed.md` says 23.

    The judge did not always quote the note whole, so a shorter needle finds
    more items: whole sentence 17, tail fragment 23, union 24 across the
    payload. Both published figures are right about what they counted. On
    office files, though, all three rules return the same 15 items and the
    same 0.0 of 29.0 — every item the choice of needle adds or removes is one
    of nine rubric items on two ``.mp4`` deliverables, which belong to the
    video-routing thread and not to this fix. A number that survives its own
    definition is worth more than a bigger one that does not."""
    payload = _stage_three_payload()
    if payload is None:
        pytest.skip("the 185-task gold-ceiling payload is not checked in here")

    by_rule = {
        "whole sentence": _disclaimer_items(payload, _FULL_SENTENCE),
        "tail fragment": _disclaimer_items(payload, _TAIL_FRAGMENT),
        "either": _disclaimer_items(payload, None),
    }
    for rule, items in by_rule.items():
        assert len(items) == MEASURED_ITEMS, rule
        assert sum(float(i["max_score"]) for _, i in items) == MEASURED_MAX_SCORE, rule
        assert sum(float(i["awarded_score"]) for _, i in items) == 0.0, rule


def test_no_item_in_that_run_recorded_what_it_actually_asked_the_tool_for():
    """Why the 15 are not attributed one by one. ``tools_used`` keeps the call
    names and drops the arguments, and ``judge_raw_response`` is null, so the
    scope that produced each empty read is not in the record. The mechanism is
    reproducible; this particular mapping is not, and no document should say
    otherwise until the payload carries the arguments."""
    payload = _stage_three_payload()
    if payload is None:
        pytest.skip("the 185-task gold-ceiling payload is not checked in here")

    items = _disclaimer_items(payload)
    assert items, "the measured set has gone missing"
    for _, item in items:
        assert item.get("judge_raw_response") is None
        assert item.get("tools_used") == ["read_deliverable"] * 3


# --------------------------------------------------------------------------
# The documents that publish these numbers must not restate them.
#
# `320-three-gaps-that-closed.md` left this thread open in writing -- "왜 같은
# 파일이 3단계에서 비어서 돌아왔는지는 이 조사에서 확인하지 않았다" -- and now
# answers it. That answer carries figures, and a figure written into prose goes
# stale silently. So these tests read the documents, re-derive every count from
# the payload above, and require the two to agree. A count that drifts in either
# direction turns the suite red instead of leaving a document quietly wrong.
# --------------------------------------------------------------------------

DOC_320 = REPO_ROOT / "tasks/rebuilding_grading_task/320-three-gaps-that-closed.md"
DOC_PR3 = REPO_ROOT / "tasks/rebuilding_grading_task/PR3_REPORT.md"

#: Sentences that were true only while the cause was unknown. The derived tests
#: below catch a number going stale; these catch a deferral being carried
#: forward after the thing it deferred was done.
RETIRED_DEFERRALS = (
    "왜 같은 파일이 3단계에서 비어서 돌아왔는지는 이 조사에서\n확인하지 않았다",
    "왜 비어서 돌아왔는지는\n   확인하지 않았고 별건으로 남긴다",
)


def _all_disclaimer_items(payload, needle):
    """Every item matching ``needle``, media included -- the totals the table
    in ``320`` publishes, before the office/media split."""
    return [
        (task, item)
        for task in payload.get("tasks") or []
        for item in task.get("items") or []
        if isinstance(item.get("evidence"), str) and needle in item["evidence"]
    ]


@pytest.mark.parametrize("doc", [DOC_320, DOC_PR3], ids=lambda p: p.name)
def test_neither_document_still_defers_the_question_it_answers(doc):
    text = doc.read_text(encoding="utf-8")
    for stale in RETIRED_DEFERRALS:
        assert stale not in text, f"{doc.name} still defers: {stale!r}"


def test_the_table_in_320_still_counts_what_the_payload_holds():
    """``320``'s table publishes 23 for the 185-task run. That figure is the
    tail-fragment count over every item, media included -- which is why this
    file's 15 does not contradict it. Both are re-derived here, so the document
    cannot keep either number after the payload stops supporting it."""
    payload = _stage_three_payload()
    if payload is None:
        pytest.skip("the 185-task gold-ceiling payload is not checked in here")

    published = _all_disclaimer_items(payload, _TAIL_FRAGMENT)
    assert f"| **0개** | 2개 | **{len(published)}개** |" in DOC_320.read_text(
        encoding="utf-8"
    ), f"320's table no longer matches the payload's {len(published)}"


def test_the_answer_320_now_gives_quotes_the_office_figures_it_can_show():
    """The prose added to ``320`` names a count, a task count and a score. All
    three come from the office subset, which is stable across every way of
    matching the sentence, and all three are re-derived here."""
    payload = _stage_three_payload()
    if payload is None:
        pytest.skip("the 185-task gold-ceiling payload is not checked in here")

    items = _disclaimer_items(payload, _FULL_SENTENCE)
    files = {f for _, i in items for f in i.get("selected_paths") or []}
    text = DOC_320.read_text(encoding="utf-8")

    awarded = sum(float(i["awarded_score"]) for _, i in items)
    maximum = sum(float(i["max_score"]) for _, i in items)
    for claim in (
        f"**{len(items)}개**가 가리키는 파일은 {len(files)}종",
        f"합계 **{awarded:.1f} / {maximum:.1f}점**",
    ):
        assert claim in text, f"320 no longer says {claim!r}"

    # The two tasks the document named as unexplained must be inside the set
    # that explains them, or the paragraph is claiming a link it does not have.
    named = {t["task_id"][:8] for t, _ in items}
    assert {"7d7fc9a7", "dfb4e0cd"} <= named


def test_the_nine_items_320_sets_aside_really_are_two_video_files():
    """``320`` sends the gap between 15 and 23 to the video-routing thread. If
    that gap ever stops being video, the sentence is wrong and the items are
    being dismissed rather than deferred."""
    payload = _stage_three_payload()
    if payload is None:
        pytest.skip("the 185-task gold-ceiling payload is not checked in here")

    union = _all_disclaimer_items(payload, _TAIL_FRAGMENT) + [
        pair
        for pair in _all_disclaimer_items(payload, _FULL_SENTENCE)
        if _TAIL_FRAGMENT not in pair[1]["evidence"]
    ]
    media = [
        (t, i)
        for t, i in union
        if any(Path(f).suffix.lower() in MEDIA_SUFFIXES
               for f in i.get("selected_paths") or [])
    ]
    assert len(media) == len(union) - MEASURED_ITEMS == 9
    names = {Path(f).name for _, i in media for f in i["selected_paths"]}
    assert names == {"GreenEnergy_v1.mp4", "Goodsin Studios CG Reel 2025.mp4"}
    for claim in sorted(names):
        assert claim in DOC_320.read_text(encoding="utf-8")


def test_every_file_the_measured_items_name_still_reads(base_dir):
    """The load-bearing claim: the files were never empty. Read all of them
    with no scope at all -- no model, no judge -- and require text from every
    one, within the range both documents publish.

    Skipped where the corpus is absent. The deliverables are a HuggingFace
    snapshot, gitignored under ``data/gdpval-local/`` and often unpacked
    elsewhere, so point ``GDPVAL_DELIVERABLE_ROOT`` at whichever copy this host
    has. There is no CI copy, so this test runs on a developer host and skips in
    CI; everything else in this file runs everywhere."""
    payload = _stage_three_payload()
    if payload is None:
        pytest.skip("the 185-task gold-ceiling payload is not checked in here")

    corpus = Path(
        os.environ.get("GDPVAL_DELIVERABLE_ROOT")
        or REPO_ROOT / "data/gdpval-local"
    )
    if not corpus.is_dir() or not any(corpus.rglob("*.xlsx")):
        pytest.skip(
            "no deliverable snapshot on this host; "
            "set GDPVAL_DELIVERABLE_ROOT to run this"
        )

    items = _disclaimer_items(payload, _FULL_SENTENCE)
    wanted = sorted({f for _, i in items for f in i.get("selected_paths") or []})
    found = {}
    for rel in wanted:
        hits = list(corpus.rglob(Path(rel).name))
        assert hits, f"{rel} is named by a graded item and is not on disk"
        response = read_deliverable(
            "read_content", hits[0].name, base_dir=str(hits[0].parent), scope={}
        )
        assert response["ok"] is True, response
        found[rel] = len(response["data"]["text"])

    assert all(n > 0 for n in found.values()), found
    text = DOC_320.read_text(encoding="utf-8")
    assert f"적게는 {min(found.values()):,}자" in text
    assert f"많게는 {max(found.values()):,}자" in text
