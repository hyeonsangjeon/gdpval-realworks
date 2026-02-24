"""Tests for core/file_preview.py"""

import os
import csv
import json
import pytest
import tempfile
from pathlib import Path

from core.file_preview import (
    generate_file_preview,
    generate_all_previews,
    _fmt_size,
)


# ─── Fixtures ───


@pytest.fixture
def tmp_dir():
    """Create temp directory for test files."""
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


@pytest.fixture
def excel_file(tmp_dir):
    """Create a sample Excel file."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sales"
    ws.append(["Product", "Revenue", "Units"])
    ws.append(["Widget A", 15000, 120])
    ws.append(["Widget B", 23000, 340])
    ws.append(["Widget C", 8500, 90])

    # Add second sheet
    ws2 = wb.create_sheet("Summary")
    ws2.append(["Metric", "Value"])
    ws2.append(["Total Revenue", 46500])

    path = tmp_dir / "sales_data.xlsx"
    wb.save(path)
    return path


@pytest.fixture
def csv_file(tmp_dir):
    """Create a sample CSV file."""
    path = tmp_dir / "data.csv"
    with open(path, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Name", "Age", "City"])
        writer.writerow(["Alice", "30", "Seoul"])
        writer.writerow(["Bob", "25", "Busan"])
    return path


@pytest.fixture
def docx_file(tmp_dir):
    """Create a sample Word document."""
    from docx import Document

    doc = Document()
    doc.add_heading("Test Report", 0)
    doc.add_paragraph("This is the introduction paragraph.")
    doc.add_paragraph("Second paragraph with more details.")
    path = tmp_dir / "report.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def pdf_file(tmp_dir):
    """Create a simple PDF file for testing."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet

        path = tmp_dir / "test.pdf"
        doc = SimpleDocTemplate(str(path), pagesize=letter)
        styles = getSampleStyleSheet()
        doc.build([
            Paragraph("Test PDF Content", styles["Title"]),
            Paragraph("Page one text body.", styles["Normal"]),
        ])
        return path
    except ImportError:
        pytest.skip("reportlab not installed")


@pytest.fixture
def text_file(tmp_dir):
    """Create a sample text file."""
    path = tmp_dir / "notes.txt"
    path.write_text("Line 1: Hello World\nLine 2: Test content\nLine 3: More text")
    return path


@pytest.fixture
def json_file(tmp_dir):
    """Create a sample JSON file."""
    path = tmp_dir / "config.json"
    path.write_text(json.dumps({"key": "value", "count": 42}, indent=2))
    return path


@pytest.fixture
def image_file(tmp_dir):
    """Create a sample PNG image."""
    try:
        from PIL import Image

        img = Image.new("RGB", (200, 100), color="red")
        path = tmp_dir / "chart.png"
        img.save(path)
        return path
    except ImportError:
        pytest.skip("Pillow not installed")


# ─── Unit Tests: _fmt_size ───


class TestFmtSize:
    def test_bytes(self):
        assert _fmt_size(500) == "500 B"

    def test_kilobytes(self):
        assert _fmt_size(2048) == "2.0 KB"

    def test_megabytes(self):
        assert _fmt_size(5 * 1024 * 1024) == "5.0 MB"


# ─── Unit Tests: generate_file_preview ───


class TestExcelPreview:
    def test_basic_preview(self, excel_file):
        preview = generate_file_preview(str(excel_file))
        assert "sales_data.xlsx" in preview
        assert "Sales" in preview
        assert "Product" in preview
        assert "Revenue" in preview
        assert "Units" in preview
        assert "Widget A" in preview
        assert "15000" in preview

    def test_shows_multiple_sheets(self, excel_file):
        preview = generate_file_preview(str(excel_file))
        assert "Sales" in preview
        assert "Summary" in preview
        assert "Total Revenue" in preview

    def test_shows_row_count(self, excel_file):
        preview = generate_file_preview(str(excel_file))
        # Sales sheet has 3 data rows
        assert "3 rows" in preview

    def test_shows_column_count(self, excel_file):
        preview = generate_file_preview(str(excel_file))
        assert "3 cols" in preview


class TestCsvPreview:
    def test_basic_preview(self, csv_file):
        preview = generate_file_preview(str(csv_file))
        assert "data.csv" in preview
        assert "Name" in preview
        assert "Age" in preview
        assert "Alice" in preview
        assert "Seoul" in preview

    def test_shows_columns(self, csv_file):
        preview = generate_file_preview(str(csv_file))
        assert "Columns:" in preview


class TestDocxPreview:
    def test_basic_preview(self, docx_file):
        preview = generate_file_preview(str(docx_file))
        assert "report.docx" in preview
        assert "introduction paragraph" in preview

    def test_shows_paragraph_count(self, docx_file):
        preview = generate_file_preview(str(docx_file))
        assert "paragraphs" in preview


class TestPdfPreview:
    def test_basic_preview(self, pdf_file):
        preview = generate_file_preview(str(pdf_file))
        assert "test.pdf" in preview
        assert "Pages:" in preview


class TestTextPreview:
    def test_basic_preview(self, text_file):
        preview = generate_file_preview(str(text_file))
        assert "notes.txt" in preview
        assert "Hello World" in preview

    def test_json_preview(self, json_file):
        preview = generate_file_preview(str(json_file))
        assert "config.json" in preview
        assert "value" in preview


class TestImagePreview:
    def test_basic_preview(self, image_file):
        preview = generate_file_preview(str(image_file))
        assert "chart.png" in preview
        assert "200" in preview
        assert "100" in preview
        assert "Dimensions" in preview


class TestMissingFile:
    def test_file_not_found(self):
        preview = generate_file_preview("/nonexistent/file.xlsx")
        assert "not found" in preview.lower()


class TestUnsupportedFile:
    def test_unsupported_extension(self, tmp_dir):
        path = tmp_dir / "model.step"
        path.write_bytes(b"\x00\x01\x02")
        preview = generate_file_preview(str(path))
        assert "model.step" in preview
        assert "not supported" in preview


# ─── Unit Tests: generate_all_previews ───


class TestGenerateAllPreviews:
    def test_empty_list(self):
        result = generate_all_previews([])
        assert result is None

    def test_single_file(self, excel_file):
        result = generate_all_previews([str(excel_file)])
        assert result is not None
        assert "REFERENCE FILES PREVIEW" in result
        assert "EXACT column names" in result
        assert "sales_data.xlsx" in result
        assert "END REFERENCE FILES" in result

    def test_multiple_files(self, excel_file, text_file):
        result = generate_all_previews([str(excel_file), str(text_file)])
        assert "sales_data.xlsx" in result
        assert "notes.txt" in result

    def test_truncation(self, excel_file, text_file):
        """Verify total output respects max_total_chars."""
        result = generate_all_previews(
            [str(excel_file), str(text_file)],
            max_total_chars=200,
        )
        assert result is not None
        assert len(result) <= 500  # header + truncated content


class TestMaxRowsLimit:
    def test_respects_max_rows(self, tmp_dir):
        """Excel preview should only show max_rows data rows."""
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["ID", "Value"])
        for i in range(50):
            ws.append([i, f"val_{i}"])

        path = tmp_dir / "big.xlsx"
        wb.save(path)

        preview = generate_file_preview(str(path), max_rows=5)
        assert "more rows" in preview
        # Should show [1] through [5], but not [6]
        assert "[5]" in preview
        assert "[6]" not in preview
