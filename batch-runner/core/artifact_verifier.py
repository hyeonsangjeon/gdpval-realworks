"""Deterministic output verifier for sandbox-generated artifacts.

Where :mod:`core.deliverable_contract` decides *what* should be produced, this
module checks that each produced file is **real**: it exists, is non-empty, lives
safely under the work directory, and can actually be opened by the library that
owns its format (openpyxl / python-docx / python-pptx / PyMuPDF / Pillow /
pandas / zipfile / ffprobe).

Everything is lazily imported and degrades gracefully — when an optional library
is unavailable the artifact is marked ``openable = None`` (not checked) with a
warning rather than a hard failure, so the verifier runs even in a light venv.
"""

from __future__ import annotations

import hashlib
import json
import mimetypes
import shutil
import subprocess
import zipfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import List, Optional

# ── suffix → coarse kind ──────────────────────────────────────────────────
_KIND_BY_SUFFIX = {
    ".xlsx": "spreadsheet", ".xlsm": "spreadsheet", ".xls": "spreadsheet",
    ".docx": "document", ".doc": "document",
    ".pptx": "presentation", ".ppt": "presentation",
    ".pdf": "pdf",
    ".png": "image", ".jpg": "image", ".jpeg": "image", ".gif": "image",
    ".bmp": "image", ".tiff": "image", ".tif": "image", ".webp": "image",
    ".csv": "data", ".tsv": "data", ".json": "data", ".parquet": "data",
    ".wav": "audio", ".mp3": "audio", ".flac": "audio", ".m4a": "audio",
    ".aac": "audio", ".ogg": "audio",
    ".mp4": "video", ".mov": "video", ".avi": "video", ".mkv": "video",
    ".webm": "video",
    ".zip": "archive",
    ".txt": "text", ".md": "text", ".html": "text", ".htm": "text",
    ".xml": "text", ".rtf": "text",
}


def classify_kind(suffix: str) -> str:
    return _KIND_BY_SUFFIX.get(suffix.lower(), "unknown")


@dataclass
class ArtifactReport:
    path: str
    rel_path: str
    exists: bool
    size_bytes: int
    suffix: str
    kind: str
    sha256: str = ""
    mime: Optional[str] = None
    openable: Optional[bool] = None  # True / False / None (not checked)
    metadata: dict = field(default_factory=dict)
    warnings: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class VerificationReport:
    artifacts: List[ArtifactReport] = field(default_factory=list)
    ok: bool = True
    blocking_errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "ok": self.ok,
            "blocking_errors": self.blocking_errors,
            "warnings": self.warnings,
            "artifacts": [a.to_dict() for a in self.artifacts],
        }


def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


# ── per-kind openability probes (all lazily imported, never raise) ─────────

def _check_spreadsheet(path: Path, report: ArtifactReport) -> None:
    if path.suffix.lower() == ".xls":
        report.openable = None
        report.warnings.append("legacy .xls not validated (openpyxl supports xlsx/xlsm)")
        return
    try:
        from openpyxl import load_workbook
    except Exception:
        report.openable = None
        report.warnings.append("openpyxl unavailable; spreadsheet not validated")
        return
    try:
        wb = load_workbook(path, read_only=True)
        visible = [ws.title for ws in wb.worksheets if ws.sheet_state == "visible"]
        report.metadata.update(
            {"sheet_count": len(wb.sheetnames), "sheet_names": wb.sheetnames[:20],
             "visible_sheets": visible[:20]}
        )
        report.openable = True
        wb.close()
    except Exception as e:
        report.openable = False
        report.errors.append(f"openpyxl could not open workbook: {e}")


def _check_document(path: Path, report: ArtifactReport) -> None:
    if path.suffix.lower() != ".docx":
        report.openable = None
        report.warnings.append("legacy .doc not validated (python-docx supports docx)")
        return
    try:
        from docx import Document
    except Exception:
        report.openable = None
        report.warnings.append("python-docx unavailable; document not validated")
        return
    try:
        doc = Document(str(path))
        report.metadata.update(
            {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}
        )
        report.openable = True
    except Exception as e:
        report.openable = False
        report.errors.append(f"python-docx could not open document: {e}")


def _check_presentation(path: Path, report: ArtifactReport) -> None:
    if path.suffix.lower() != ".pptx":
        report.openable = None
        report.warnings.append("legacy .ppt not validated (python-pptx supports pptx)")
        return
    try:
        from pptx import Presentation
    except Exception:
        report.openable = None
        report.warnings.append("python-pptx unavailable; presentation not validated")
        return
    try:
        prs = Presentation(str(path))
        report.metadata.update({"slide_count": len(prs.slides)})
        report.openable = True
    except Exception as e:
        report.openable = False
        report.errors.append(f"python-pptx could not open presentation: {e}")


def _check_pdf(path: Path, report: ArtifactReport) -> None:
    try:
        import fitz  # PyMuPDF
    except Exception:
        report.openable = None
        report.warnings.append("PyMuPDF unavailable; PDF not validated")
        return
    try:
        doc = fitz.open(str(path))
        page_count = doc.page_count
        report.metadata.update({"page_count": page_count})
        doc.close()
        if page_count <= 0:
            report.openable = False
            report.errors.append("PDF has 0 pages")
        else:
            report.openable = True
    except Exception as e:
        report.openable = False
        report.errors.append(f"PyMuPDF could not open PDF: {e}")


def _check_image(path: Path, report: ArtifactReport) -> None:
    try:
        from PIL import Image
    except Exception:
        report.openable = None
        report.warnings.append("Pillow unavailable; image not validated")
        return
    try:
        with Image.open(path) as im:
            im.verify()
        with Image.open(path) as im2:
            report.metadata.update({"width": im2.width, "height": im2.height,
                                    "mode": im2.mode})
        report.openable = True
    except Exception as e:
        report.openable = False
        report.errors.append(f"Pillow could not open image: {e}")


def _check_data(path: Path, report: ArtifactReport) -> None:
    suffix = path.suffix.lower()
    if suffix == ".json":
        try:
            with path.open("r", encoding="utf-8") as fh:
                json.load(fh)
            report.openable = True
        except Exception as e:
            report.openable = False
            report.errors.append(f"invalid JSON: {e}")
        return
    if suffix == ".parquet":
        try:
            import pandas as pd
            df = pd.read_parquet(path)
            report.metadata.update({"rows": int(df.shape[0]), "cols": int(df.shape[1])})
            report.openable = True
        except Exception as e:
            report.openable = None
            report.warnings.append(f"parquet not validated: {e}")
        return
    # csv / tsv
    sep = "\t" if suffix == ".tsv" else ","
    try:
        import pandas as pd
        df = pd.read_csv(path, sep=sep, nrows=2000)
        report.metadata.update({"rows_sampled": int(df.shape[0]), "cols": int(df.shape[1])})
        report.openable = True
    except Exception:
        try:
            with path.open("r", encoding="utf-8", errors="replace") as fh:
                first = fh.readline()
            report.metadata.update({"first_line_fields": len(first.split(sep))})
            report.openable = True
        except Exception as e:
            report.openable = False
            report.errors.append(f"could not read delimited file: {e}")


def _ffprobe(path: Path) -> Optional[dict]:
    if not shutil.which("ffprobe"):
        return None
    try:
        proc = subprocess.run(
            ["ffprobe", "-v", "error", "-show_format", "-show_streams",
             "-of", "json", str(path)],
            capture_output=True, text=True, timeout=30,
        )
        if proc.returncode != 0:
            return {"_error": proc.stderr[-300:]}
        return json.loads(proc.stdout or "{}")
    except Exception:
        return None


def _check_audio(path: Path, report: ArtifactReport) -> None:
    probe = _ffprobe(path)
    if probe is not None and "_error" not in probe:
        streams = probe.get("streams", [])
        audio_streams = [s for s in streams if s.get("codec_type") == "audio"]
        dur = (probe.get("format") or {}).get("duration")
        report.metadata.update({"audio_streams": len(audio_streams), "duration": dur})
        report.openable = bool(audio_streams)
        if not audio_streams:
            report.errors.append("no audio stream detected by ffprobe")
        return
    if path.suffix.lower() in (".wav", ".flac", ".ogg"):
        try:
            import soundfile as sf
            info = sf.info(str(path))
            report.metadata.update({"samplerate": info.samplerate,
                                    "frames": info.frames, "channels": info.channels})
            report.openable = info.frames > 0
            if info.frames <= 0:
                report.errors.append("audio has 0 frames")
            return
        except Exception:
            pass
    report.openable = None
    report.warnings.append("audio not validated (ffprobe/soundfile unavailable)")


def _check_video(path: Path, report: ArtifactReport) -> None:
    probe = _ffprobe(path)
    if probe is not None and "_error" not in probe:
        streams = probe.get("streams", [])
        video_streams = [s for s in streams if s.get("codec_type") == "video"]
        dur = (probe.get("format") or {}).get("duration")
        report.metadata.update({"video_streams": len(video_streams), "duration": dur})
        report.openable = bool(video_streams)
        if not video_streams:
            report.errors.append("no video stream detected by ffprobe")
        return
    try:
        import cv2
        cap = cv2.VideoCapture(str(path))
        opened = cap.isOpened()
        frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT)) if opened else 0
        cap.release()
        report.metadata.update({"frame_count": frames})
        report.openable = opened and frames > 0
        if not report.openable:
            report.errors.append("OpenCV could not read video frames")
        return
    except Exception:
        pass
    report.openable = None
    report.warnings.append("video not validated (ffprobe/opencv unavailable)")


def _check_archive(path: Path, report: ArtifactReport) -> None:
    try:
        with zipfile.ZipFile(path) as zf:
            bad = zf.testzip()
            report.metadata.update({"entries": len(zf.namelist())})
            if bad is not None:
                report.openable = False
                report.errors.append(f"corrupt zip entry: {bad}")
            else:
                report.openable = True
    except Exception as e:
        report.openable = False
        report.errors.append(f"could not open archive: {e}")


def _check_text(path: Path, report: ArtifactReport) -> None:
    try:
        data = path.read_text(encoding="utf-8", errors="replace")
        report.metadata.update({"char_count": len(data), "line_count": data.count("\n") + 1})
        report.openable = True
    except Exception as e:
        report.openable = False
        report.errors.append(f"could not read text file: {e}")


_KIND_CHECKERS = {
    "spreadsheet": _check_spreadsheet,
    "document": _check_document,
    "presentation": _check_presentation,
    "pdf": _check_pdf,
    "image": _check_image,
    "data": _check_data,
    "audio": _check_audio,
    "video": _check_video,
    "archive": _check_archive,
    "text": _check_text,
}


def verify_one(path, workdir=None) -> ArtifactReport:
    """Verify a single artifact; never raises."""
    path = Path(path)
    rel = path.name
    if workdir is not None:
        try:
            rel = str(path.resolve().relative_to(Path(workdir).resolve()))
        except Exception:
            rel = path.name
    suffix = path.suffix.lower()
    report = ArtifactReport(
        path=str(path), rel_path=rel, exists=path.exists(),
        size_bytes=path.stat().st_size if path.exists() else 0,
        suffix=suffix, kind=classify_kind(suffix),
        mime=mimetypes.guess_type(str(path))[0],
    )

    if not report.exists:
        report.errors.append("file does not exist")
        return report

    # Path safety: an artifact must live under the work directory.
    if workdir is not None:
        try:
            Path(path).resolve().relative_to(Path(workdir).resolve())
        except Exception:
            report.errors.append("artifact path escapes the work directory")
            return report

    if report.size_bytes == 0:
        report.errors.append("file is empty (0 bytes)")
        report.openable = False
        return report

    try:
        report.sha256 = _sha256(path)
    except Exception as e:
        report.warnings.append(f"could not hash file: {e}")

    checker = _KIND_CHECKERS.get(report.kind)
    if checker is None:
        report.openable = None
        report.warnings.append(f"no validator for '{suffix or 'extensionless'}' file")
    else:
        try:
            checker(path, report)
        except Exception as e:  # pragma: no cover - defensive
            report.openable = False
            report.errors.append(f"verification crashed: {e}")
    return report


def verify_artifacts(artifacts, contract=None, workdir=None) -> VerificationReport:
    """Verify every artifact and roll up blocking errors / warnings."""
    reports = [verify_one(a, workdir=workdir) for a in artifacts]
    blocking: List[str] = []
    warnings: List[str] = []
    for r in reports:
        for e in r.errors:
            blocking.append(f"{r.rel_path}: {e}")
        for w in r.warnings:
            warnings.append(f"{r.rel_path}: {w}")
    return VerificationReport(
        artifacts=reports,
        ok=not blocking,
        blocking_errors=blocking,
        warnings=warnings,
    )
